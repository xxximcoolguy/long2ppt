"""longtext-to-slides · 通用文本 + 图片提取脚本

支持格式：
- PDF  → pdfplumber (fallback: pypdf) + 图片提取
- DOCX → python-docx + 内嵌图片提取
- EPUB → ebooklib + ITEM_IMAGE
- MD / TXT → 直接读取 + 解析 ![](path) 引用
- HTML → BeautifulSoup 抽取正文 + <img> 引用

输出：
- <output.md>：UTF-8 markdown，含页/段标记，图片位置用 `![alt](assets/img-NNN.ext)` 占位
- <output_dir>/assets/img-NNN.{png,jpg,...}：提取的图片文件
- <output.json>：元信息 { format, chars, images: [{path, page, alt, size}] }

Usage:
    python extract.py <input_file> <output.md>
"""
import sys
import os
import logging
import re
import json
import hashlib
import shutil
from pathlib import Path
from urllib.parse import urlparse, unquote

logging.getLogger("pdfminer").setLevel(logging.ERROR)
logging.getLogger("pdfplumber").setLevel(logging.ERROR)


# ============================================================
# Image helpers
# ============================================================

class ImageBag:
    """收集图片，统一命名 / 去重 / 输出。"""
    def __init__(self, output_md_path: str):
        self.md_path = Path(output_md_path)
        self.assets_dir = self.md_path.parent / "assets"
        self.assets_dir.mkdir(parents=True, exist_ok=True)
        self.images: list[dict] = []
        self._seen_hashes: dict[str, str] = {}  # hash -> relative path

    def add_binary(self, data: bytes, ext: str, page: int | None = None, alt: str = "") -> str:
        """加入二进制图片，返回 markdown 引用路径。"""
        h = hashlib.md5(data).hexdigest()[:10]
        if h in self._seen_hashes:
            return self._seen_hashes[h]

        idx = len(self.images) + 1
        ext = ext.lstrip(".").lower()
        if ext not in ("png", "jpg", "jpeg", "gif", "webp", "svg", "bmp"):
            ext = "png"
        filename = f"img-{idx:03d}.{ext}"
        out_path = self.assets_dir / filename
        out_path.write_bytes(data)

        rel = f"assets/{filename}"
        self._seen_hashes[h] = rel
        self.images.append({
            "rel": rel,
            "size": len(data),
            "page": page,
            "alt": alt,
        })
        return rel

    def add_from_path(self, src: str, base: Path, page: int | None = None, alt: str = "") -> str | None:
        """从本地路径或 URL 加入图片。"""
        # 远程 URL
        if src.startswith(("http://", "https://")):
            try:
                import urllib.request
                with urllib.request.urlopen(src, timeout=10) as resp:
                    data = resp.read()
                ext = Path(urlparse(src).path).suffix or ".png"
                return self.add_binary(data, ext, page=page, alt=alt)
            except Exception as e:
                print(f"[!] failed to download {src}: {e}", file=sys.stderr)
                return None

        # 本地路径
        src_path = Path(unquote(src))
        if not src_path.is_absolute():
            src_path = (base / src_path).resolve()
        if not src_path.exists():
            print(f"[!] image not found: {src_path}", file=sys.stderr)
            return None
        try:
            data = src_path.read_bytes()
            return self.add_binary(data, src_path.suffix, page=page, alt=alt)
        except Exception as e:
            print(f"[!] failed to read {src_path}: {e}", file=sys.stderr)
            return None

    def to_meta(self) -> list[dict]:
        return self.images


# ============================================================
# PDF
# ============================================================

def extract_pdf(pdf_path: str, bag: ImageBag) -> str:
    try:
        import pdfplumber
    except ImportError:
        return _extract_pdf_pypdf(pdf_path, bag)

    parts: list[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            parts.append(f"\n\n<!-- ===== Page {i} ===== -->\n\n{text}")

            # 图片
            for j, img in enumerate(page.images, start=1):
                try:
                    # 用 page crop 截取图片区域
                    x0, top, x1, bottom = img["x0"], img["top"], img["x1"], img["bottom"]
                    cropped = page.crop((x0, top, x1, bottom))
                    pil = cropped.to_image(resolution=144)
                    import io
                    buf = io.BytesIO()
                    pil.save(buf, format="PNG")
                    rel = bag.add_binary(buf.getvalue(), "png", page=i, alt=f"Page {i} Image {j}")
                    parts.append(f"\n\n![Page {i} Image {j}]({rel})\n")
                except Exception as e:
                    print(f"[!] page {i} img {j} extract failed: {e}", file=sys.stderr)
    return "\n".join(parts)


def _extract_pdf_pypdf(pdf_path: str, bag: ImageBag) -> str:
    from pypdf import PdfReader
    reader = PdfReader(pdf_path)
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        parts.append(f"\n\n<!-- ===== Page {i} ===== -->\n\n{text}")
        # pypdf 图片提取
        try:
            for j, img in enumerate(page.images, start=1):
                rel = bag.add_binary(img.data, img.name.rsplit(".", 1)[-1] if "." in img.name else "png",
                                     page=i, alt=img.name)
                parts.append(f"\n\n![{img.name}]({rel})\n")
        except Exception as e:
            print(f"[!] page {i} image extract failed: {e}", file=sys.stderr)
    return "\n".join(parts)


# ============================================================
# DOCX
# ============================================================

def extract_docx(docx_path: str, bag: ImageBag) -> str:
    from docx import Document
    doc = Document(docx_path)
    parts: list[str] = []

    # 把 inline image 的 relationship id → bag 的 rel 路径
    image_map: dict[str, str] = {}
    for rid, rel in doc.part.rels.items():
        if "image" in rel.reltype.lower():
            try:
                blob = rel.target_part.blob
                ext = Path(rel.target_part.partname).suffix or ".png"
                image_map[rid] = bag.add_binary(blob, ext, alt=rel.target_part.partname)
            except Exception as e:
                print(f"[!] docx image rid={rid} failed: {e}", file=sys.stderr)

    # 遍历段落 / 处理 inline image
    from docx.oxml.ns import qn
    for p in doc.paragraphs:
        style = p.style.name if p.style else ""
        text = p.text.strip()
        # 检查段落里是否有 inline image
        inline_imgs = []
        for run in p.runs:
            for blip in run._element.findall(".//" + qn("a:blip")):
                rid = blip.get(qn("r:embed"))
                if rid and rid in image_map:
                    inline_imgs.append(image_map[rid])

        if inline_imgs and not text:
            for rel in inline_imgs:
                parts.append(f"\n\n![]({rel})\n")
            continue

        if not text:
            parts.append("")
            continue

        if "Heading 1" in style:
            parts.append(f"# {text}")
        elif "Heading 2" in style:
            parts.append(f"## {text}")
        elif "Heading 3" in style:
            parts.append(f"### {text}")
        else:
            parts.append(text)

        # 段后追加 inline img
        for rel in inline_imgs:
            parts.append(f"\n![]({rel})\n")

    # 表格
    for table in doc.tables:
        parts.append("\n\n<!-- TABLE -->\n")
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append("| " + " | ".join(cells) + " |")

    return "\n".join(parts)


# ============================================================
# EPUB
# ============================================================

def extract_epub(epub_path: str, bag: ImageBag) -> str:
    from ebooklib import epub, ITEM_DOCUMENT, ITEM_IMAGE
    from bs4 import BeautifulSoup

    book = epub.read_epub(epub_path)

    # 图片资源 → bag
    image_map: dict[str, str] = {}
    for item in book.get_items_of_type(ITEM_IMAGE):
        try:
            ext = Path(item.file_name).suffix or ".png"
            rel = bag.add_binary(item.get_content(), ext, alt=item.file_name)
            image_map[item.file_name] = rel
            # 也保留只取 basename 的索引
            image_map[Path(item.file_name).name] = rel
        except Exception as e:
            print(f"[!] epub image {item.file_name} failed: {e}", file=sys.stderr)

    parts: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        # 把 <img> 转 markdown 图片
        for img in soup.find_all("img"):
            src = img.get("src", "")
            alt = img.get("alt", "")
            # epub 内部相对路径
            key = src
            if key not in image_map:
                key = Path(src).name
            if key in image_map:
                img.replace_with(f"\n\n![{alt}]({image_map[key]})\n\n")
        text = soup.get_text("\n", strip=True)
        if text:
            parts.append(f"\n\n<!-- ===== Chapter ===== -->\n\n{text}")
    return "\n".join(parts)


# ============================================================
# HTML
# ============================================================

def extract_html(html_path: str, bag: ImageBag) -> str:
    from bs4 import BeautifulSoup
    base = Path(html_path).parent
    with open(html_path, "r", encoding="utf-8") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()

    # 处理 <img>
    for img in soup.find_all("img"):
        src = img.get("src", "")
        alt = img.get("alt", "")
        if src:
            rel = bag.add_from_path(src, base, alt=alt)
            if rel:
                img.replace_with(f"\n\n![{alt}]({rel})\n\n")

    return soup.get_text("\n", strip=True)


# ============================================================
# Markdown / TXT
# ============================================================

MD_IMG_PATTERN = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")


# ============================================================
# Language detection
# ============================================================

def detect_language(text: str) -> dict:
    """返回 { dominant_lang, han_chars, latin_chars, han_ratio, needs_translation }

    - han_ratio：中文字符占（中文+英文字母）的比例
    - dominant_lang：'zh' (han_ratio >= 0.5) / 'en' (han_ratio < 0.1) / 'bilingual' (其他)
    - needs_translation：dominant == 'en' 且 han_ratio < 0.05 时为 True
    """
    han = sum(1 for c in text if '一' <= c <= '鿿')
    latin = sum(1 for c in text if c.isalpha() and c.isascii())
    total = han + latin
    if total == 0:
        return {
            "dominant_lang": "unknown",
            "han_chars": 0, "latin_chars": 0,
            "han_ratio": 0.0, "needs_translation": False,
        }
    ratio = han / total
    if ratio >= 0.55:
        dominant = "zh"
    elif ratio < 0.10:
        dominant = "en"
    else:
        dominant = "bilingual"
    return {
        "dominant_lang": dominant,
        "han_chars": han,
        "latin_chars": latin,
        "han_ratio": round(ratio, 3),
        "needs_translation": dominant == "en" and ratio < 0.05,
    }



def extract_markdown(md_path: str, bag: ImageBag) -> str:
    """Markdown：原样读取，替换其中的 ![](path) 引用为 bag 输出路径。"""
    base = Path(md_path).parent
    with open(md_path, "r", encoding="utf-8") as f:
        text = f.read()

    def repl(m: re.Match) -> str:
        alt, src = m.group(1), m.group(2)
        # 已经指向 assets/ 的就保留
        if src.startswith("assets/"):
            return m.group(0)
        rel = bag.add_from_path(src, base, alt=alt)
        if rel:
            return f"![{alt}]({rel})"
        return m.group(0)  # 保留原 markdown，让阅读器处理

    return MD_IMG_PATTERN.sub(repl, text)


def extract_text(path: str, bag: ImageBag) -> str:
    """TXT：无图片，直接读取。"""
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


# ============================================================
# Dispatch
# ============================================================

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".epub": extract_epub,
    ".html": extract_html,
    ".htm": extract_html,
    ".md": extract_markdown,
    ".markdown": extract_markdown,
    ".txt": extract_text,
}


def extract(input_path: str, output_path: str) -> dict:
    p = Path(input_path)
    ext = p.suffix.lower()
    if ext not in EXTRACTORS:
        raise SystemExit(f"Unsupported file type: {ext}")

    extractor = EXTRACTORS[ext]
    bag = ImageBag(output_path)

    print(f"[*] Extracting {p.name} via {extractor.__name__}...", file=sys.stderr)

    try:
        text = extractor(str(p), bag)
    except ImportError as e:
        sys.exit(f"Missing dependency for {ext}: {e}. "
                 f"Install: pdfplumber / python-docx / ebooklib / beautifulsoup4 as needed.")

    # 清理多余空行
    text = re.sub(r"\n{4,}", "\n\n\n", text)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(text, encoding="utf-8")

    # 语言检测：中文字符 vs 英文字母比例
    lang_info = detect_language(text)

    images = bag.to_meta()
    meta_path = out.with_suffix(".meta.json")
    meta = {
        "format": ext,
        "chars": len(text.replace(" ", "").replace("\n", "")),
        "lines": text.count("\n"),
        "output": str(out),
        "images": images,
        "image_count": len(images),
        "image_total_bytes": sum(i["size"] for i in images),
        "assets_dir": str(bag.assets_dir),
        **lang_info,
    }
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")
    return meta


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python extract.py <input_file> <output.md>", file=sys.stderr)
        sys.exit(1)

    result = extract(sys.argv[1], sys.argv[2])
    print(f"OK  format={result['format']}  chars={result['chars']}  "
          f"images={result['image_count']}  bytes={result['image_total_bytes']}  "
          f"out={result['output']}")
